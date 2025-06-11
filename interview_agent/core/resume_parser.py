"""
简历解析模块 - 支持多种文件格式和信息抽取
"""

import json
import re
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
import os

# 文档解析相关
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import markdown
    from bs4 import BeautifulSoup
except ImportError:
    markdown = None
    BeautifulSoup = None

try:
    import docx
except ImportError:
    docx = None

# LLM相关
from .llm_client import Message
from config.settings import settings


@dataclass
class CandidateProfile:
    """候选人简历信息"""
    name: str
    experience_years: int = 0
    skills: List[str] = field(default_factory=list)
    education: List[str] = field(default_factory=list)
    work_experience: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    contact_info: Dict[str, str] = field(default_factory=dict)
    languages: List[str] = field(default_factory=list)
    summary: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式"""
        return {
            "name": self.name,
            "experience_years": self.experience_years,
            "skills": self.skills,
            "education": self.education,
            "work_experience": self.work_experience,
            "projects": self.projects,
            "certifications": self.certifications,
            "contact_info": self.contact_info,
            "languages": self.languages,
            "summary": self.summary
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'CandidateProfile':
        """从字典创建实例"""
        return cls(
            name=data.get("name", ""),
            experience_years=data.get("experience_years", 0),
            skills=data.get("skills", []),
            education=data.get("education", []),
            work_experience=data.get("work_experience", []),
            projects=data.get("projects", []),
            certifications=data.get("certifications", []),
            contact_info=data.get("contact_info", {}),
            languages=data.get("languages", []),
            summary=data.get("summary", "")
        )


@dataclass
class ParsedDocument:
    """解析后的文档数据结构"""
    raw_text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    structured_data: Dict[str, Any] = field(default_factory=dict)
    file_type: str = ""
    file_path: Optional[Path] = None


class DocumentParser(ABC):
    """文档解析器基类"""
    
    @abstractmethod
    def can_parse(self, file_path: Path) -> bool:
        """检查是否可以解析该文件"""
        pass
    
    @abstractmethod
    def parse(self, file_path: Path) -> ParsedDocument:
        """解析文档"""
        pass


class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """使用pypdf进行高级PDF解析"""
        if pypdf is None:
            raise ImportError("无法解析PDF文件: 缺少pypdf库。请使用'pip install pypdf'安装。")
            
        text_content = []
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # 提取元数据
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    }
                
                # 提取文本内容
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                
                raw_text = '\n'.join(text_content)
                
                return ParsedDocument(
                    raw_text=raw_text,
                    metadata=metadata,
                    file_type='pdf',
                    file_path=file_path
                )
                
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")
    

class MarkdownParser(DocumentParser):
    """Markdown文档解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.md', '.markdown']
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """解析Markdown文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 使用markdown库解析
            html_content = markdown.markdown(content)
            
            # 提取纯文本（简单实现，可以根据需要增强）
            text_content = re.sub('<[^<]+?>', '', html_content)
            
            return ParsedDocument(
                raw_text=content,  # 保留原始markdown格式
                metadata={'format': 'markdown'},
                file_type='markdown',
                file_path=file_path
            )
            
        except Exception as e:
            raise Exception(f"Markdown解析失败: {str(e)}")


class TextParser(DocumentParser):
    """纯文本文档解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.text']
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """解析纯文本文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return ParsedDocument(
                raw_text=content,
                metadata={'format': 'text'},
                file_type='text',
                file_path=file_path
            )
            
        except Exception as e:
            raise Exception(f"文本解析失败: {str(e)}")


class DocxParser(DocumentParser):
    """Word文档解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def parse(self, file_path: Path) -> ParsedDocument:
        """解析Word文档"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # 提取所有段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            raw_text = '\n'.join(text_content)
            
            return ParsedDocument(
                raw_text=raw_text,
                metadata={'format': 'docx'},
                file_type='docx',
                file_path=file_path
            )
            
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}") 


class UniversalDocumentParser:
    """通用文档解析器 - 支持多种格式"""
    
    def __init__(self):
        self.parsers = [
            PDFParser(),
            MarkdownParser(),
            TextParser(),
            DocxParser(),
        ]
    
    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """解析文档，自动识别格式"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 找到合适的解析器
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser.parse(file_path)
        
        # 如果没有合适的解析器，尝试作为文本文件处理
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return ParsedDocument(
                raw_text=content,
                metadata={'format': 'unknown'},
                file_type='unknown',
                file_path=file_path
            )
        except Exception as e:
            raise Exception(f"不支持的文件格式或解析失败: {str(e)}")


class LLMExtractor:
    """基于LLM的信息抽取器"""
    
    def __init__(self, llm_client=None):
        self.llm = llm_client
        # 从settings获取max_tokens，如果未设置则使用默认值4000
        self.max_tokens = getattr(settings, 'extractor_max_tokens', 4000)
    
    def extract_structured_info(self, 
                              document: ParsedDocument,
                              extraction_schema: Dict[str, Any],
                              additional_instructions: str = "") -> Dict[str, Any]:
        """
        使用LLM从文档中抽取结构化信息
        
        Args:
            document: 解析后的文档
            extraction_schema: 期望抽取的信息结构
            additional_instructions: 额外的抽取指令
        
        Returns:
            抽取的结构化信息
        """
        
        if not self.llm:
            raise ValueError("LLM客户端未设置，请确保在使用前设置LLM客户端")
            
        system_prompt = """你是一个专业的文档信息抽取助手。
请根据提供的文档内容和抽取模式，准确地提取相关信息。

注意事项：
1. 严格按照提供的模式结构返回JSON格式数据
2. 如果某个字段在文档中找不到，使用null或空值
3. 保持信息的准确性，不要编造内容
4. 对于列表类型的字段，如果没有找到任何项，返回空列表[]
5. 日期格式统一为 YYYY-MM-DD
6. 保持专业术语的准确性"""

        user_prompt = f"""请从以下文档中抽取信息：

【文档内容】
{document.raw_text[:4000]}  # 限制长度，避免超出token限制

【期望抽取的信息结构】
{json.dumps(extraction_schema, ensure_ascii=False, indent=2)}

{f'【额外说明】{additional_instructions}' if additional_instructions else ''}

请严格按照上述结构返回JSON格式的抽取结果。"""

        messages = [
            Message(role="system", content=system_prompt),
            Message(role="user", content=user_prompt)
        ]
        
        try:
            response = self.llm.chat_completion(
                messages,
                temperature=0.1,  # 低温度以提高准确性
                max_tokens=self.max_tokens
            )
            
            # 尝试解析JSON响应
            content = response.content
            
            # 清理可能的markdown代码块标记
            if '```json' in content:
                content = content.split('```json')[1].split('```')[0]
            elif '```' in content:
                content = content.split('```')[1].split('```')[0]
            
            extracted_data = json.loads(content.strip())
            return extracted_data
            
        except json.JSONDecodeError:
            # 如果JSON解析失败，尝试使用正则提取
            return self._fallback_extraction(response.content, extraction_schema)
        except Exception as e:
            raise Exception(f"LLM信息抽取失败: {str(e)}")
    
    def _fallback_extraction(self, 
                           llm_response: str, 
                           schema: Dict[str, Any]) -> Dict[str, Any]:
        """降级的抽取方法"""
        # 创建一个基于schema的空结果
        result = {}
        for key, value in schema.items():
            if isinstance(value, list):
                result[key] = []
            elif isinstance(value, dict):
                result[key] = {}
            else:
                result[key] = None
        
        # 这里可以添加一些基础的正则提取逻辑
        # 比如提取邮箱、电话等
        
        return result


class ResumeParser:
    """简历解析器 - 整合文档解析和信息抽取"""
    
    def __init__(self, 
                 extraction_schema: Optional[Dict[str, Any]] = None):
        """
        初始化简历解析器
        
        Args:
            extraction_schema: 自定义的信息抽取模式，如果不提供则使用默认模式
        """
        self.document_parser = UniversalDocumentParser()
        self.extractor = LLMExtractor()
        
        # 默认的简历信息抽取模式
        self.default_schema = {
            "basic_info": {
                "name": "string",
                "email": "string",
                "phone": "string",
                "location": "string",
                "summary": "string"
            },
            "education": [
                {
                    "school": "string",
                    "degree": "string",
                    "major": "string",
                    "start_date": "string",
                    "end_date": "string",
                    "gpa": "string"
                }
            ],
            "work_experience": [
                {
                    "company": "string",
                    "position": "string",
                    "start_date": "string",
                    "end_date": "string",
                    "description": "string",
                    "achievements": ["string"]
                }
            ],
            "projects": [
                {
                    "name": "string",
                    "role": "string",
                    "description": "string",
                    "technologies": ["string"],
                    "achievements": ["string"]
                }
            ],
            "skills": {
                "technical": ["string"],
                "languages": ["string"],
                "tools": ["string"],
                "soft_skills": ["string"]
            },
            "certifications": [
                {
                    "name": "string",
                    "issuer": "string",
                    "date": "string"
                }
            ]
        }
        
        self.extraction_schema = extraction_schema or self.default_schema
    
    def parse(self, 
              file_path: Union[str, Path],
              custom_schema: Optional[Dict[str, Any]] = None,
              additional_instructions: str = "") -> Dict[str, Any]:
        """
        解析简历文件并抽取结构化信息
        
        Args:
            file_path: 简历文件路径
            custom_schema: 自定义抽取模式（可选）
            additional_instructions: 额外的抽取指令（可选）
        
        Returns:
            包含解析文本和结构化信息的字典
        """
        # 1. 解析文档
        document = self.document_parser.parse(file_path)
        
        # 2. 使用LLM抽取信息
        schema = custom_schema or self.extraction_schema
        structured_info = self.extractor.extract_structured_info(
            document,
            schema,
            additional_instructions
        )
        
        # 3. 返回完整结果
        return {
            "raw_text": document.raw_text,
            "metadata": document.metadata,
            "file_type": document.file_type,
            "structured_info": structured_info
        }
    
    def parse_to_profile(self,
                       file_path: Union[str, Path],
                       custom_schema: Optional[Dict[str, Any]] = None,
                       additional_instructions: str = "") -> CandidateProfile:
        """
        解析简历文件并直接返回CandidateProfile对象
        
        Args:
            file_path: 简历文件路径
            custom_schema: 自定义抽取模式（可选）
            additional_instructions: 额外的抽取指令（可选）
            
        Returns:
            CandidateProfile对象
        """
        # 解析简历
        result = self.parse(file_path, custom_schema, additional_instructions)
        structured_info = result["structured_info"]
        
        # 提取基本信息
        basic_info = structured_info.get("basic_info", {})
        name = basic_info.get("name", "未知")
        summary = basic_info.get("summary", "")
        
        # 提取技能
        skills_info = structured_info.get("skills", {})
        skills = []
        skills.extend(skills_info.get("technical", []))
        skills.extend(skills_info.get("tools", []))
        
        # 提取教育经历
        education_list = structured_info.get("education", [])
        education = []
        for edu in education_list:
            edu_str = f"{edu.get('school', '')} - {edu.get('degree', '')} {edu.get('major', '')}"
            if edu.get('start_date') and edu.get('end_date'):
                edu_str += f" ({edu.get('start_date')} - {edu.get('end_date')})"
            education.append(edu_str)
        
        # 提取工作经历
        work_list = structured_info.get("work_experience", [])
        work_experience = []
        experience_years = 0
        for work in work_list:
            work_str = f"{work.get('company', '')} - {work.get('position', '')}"
            if work.get('start_date') and work.get('end_date'):
                work_str += f" ({work.get('start_date')} - {work.get('end_date')})"
                # 尝试计算工作年限
                try:
                    start_year = int(work.get('start_date', '').split('-')[0])
                    end_year = int(work.get('end_date', '').split('-')[0])
                    experience_years += (end_year - start_year)
                except (ValueError, IndexError):
                    pass
            if work.get('description'):
                work_str += f": {work.get('description')}"
            work_experience.append(work_str)
        
        # 提取项目经历
        project_list = structured_info.get("projects", [])
        projects = []
        for proj in project_list:
            proj_str = f"{proj.get('name', '')} - {proj.get('role', '')}"
            if proj.get('description'):
                proj_str += f": {proj.get('description')}"
            if proj.get('technologies'):
                proj_str += f" [技术: {', '.join(proj.get('technologies', []))}]"
            projects.append(proj_str)
        
        # 提取证书
        cert_list = structured_info.get("certifications", [])
        certifications = []
        for cert in cert_list:
            cert_str = f"{cert.get('name', '')} - {cert.get('issuer', '')}"
            if cert.get('date'):
                cert_str += f" ({cert.get('date')})"
            certifications.append(cert_str)
        
        # 提取语言
        languages = skills_info.get("languages", [])
        
        # 提取联系信息
        contact_info = {
            "email": basic_info.get("email", ""),
            "phone": basic_info.get("phone", ""),
            "location": basic_info.get("location", "")
        }
        
        # 创建并返回CandidateProfile对象
        return CandidateProfile(
            name=name,
            experience_years=experience_years,
            skills=skills,
            education=education,
            work_experience=work_experience,
            projects=projects,
            certifications=certifications,
            contact_info=contact_info,
            languages=languages,
            summary=summary
        )
    
    def parse_batch(self, 
                   file_paths: List[Union[str, Path]],
                   custom_schema: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """批量解析简历文件"""
        results = []
        for file_path in file_paths:
            try:
                result = self.parse(file_path, custom_schema)
                result["file_path"] = str(file_path)
                result["status"] = "success"
                results.append(result)
            except Exception as e:
                results.append({
                    "file_path": str(file_path),
                    "status": "error",
                    "error": str(e)
                })
        return results 